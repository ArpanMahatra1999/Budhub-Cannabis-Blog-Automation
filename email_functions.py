# email and document libraries
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
from docx import Document


def create_document(headline, description):
    doc = Document()
    doc.add_heading(headline[1:len(headline)-1], level=1)
    doc.add_paragraph(description[1:len(description)-1])
    filename = "blog.docx"
    doc.save(filename)
    return filename


def send_email_with_attachment(sender, password, receiver, subject, body, headline, description):
    # create email message
    msg = MIMEMultipart()
    msg["From"] = sender
    msg["To"] = receiver
    msg["Subject"] = subject
    msg.attach(MIMEText(body, "plain"))

    # Create document and attach
    filename = create_document(headline, description)
    with open(filename, "rb") as attachment:
        part = MIMEBase("application", "octet-stream")
        part.set_payload(attachment.read())
        encoders.encode_base64(part)
        part.add_header("Content-Disposition", f"attachment; filename={filename}")
        msg.attach(part)

    # Send the email
    with smtplib.SMTP("smtp.gmail.com", 587) as server:
        server.starttls()
        server.login(sender, password)
        server.send_message(msg)
        print("Email sent with attachment!")