from docx import Document
from datetime import datetime


def create_doc(headline, description):
    """
    Creates a Word document with a headline and description,
    and saves it with a filename based on the current date and time.
    """
    doc = Document()
    doc.add_heading(headline, level=1)
    doc.add_paragraph(description)
    date = datetime.now().strftime("%Y-%m-%d")
    time = datetime.now().strftime("%H-%M-%S")
    filename = f"blogs/blog_{date}-{time}.docx"

    doc.save(filename)
    print(f"Document saved as: {filename}")
